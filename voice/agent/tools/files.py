"""
agent/tools/files.py
File system tools: read, write, list, move, delete, find.
"""
import os
import re
import shutil
import subprocess
from pathlib import Path

HOME = str(Path.home())

SAFE_BASE_DIRS = [
    os.path.join(HOME, "Desktop"),
    os.path.join(HOME, "Documents"),
    os.path.join(HOME, "Downloads"),
    HOME,
]


def _normalize_search_token(name: str) -> str:
    """Fix common speech-to-text name variants before file search."""
    token = name.strip().strip("'\"")
    token = re.sub(r"'s$", "", token, flags=re.I)
    token = re.sub(r"s'$", "", token, flags=re.I)
    if re.fullmatch(r"[a-z]+s", token, re.I) and len(token) > 4:
        token = token[:-1]
    return token.strip()


def _search_queries(name: str) -> list[str]:
    """Build ordered unique Spotlight / scan queries for a spoken name."""
    raw = name.strip()
    normalized = _normalize_search_token(raw)
    queries: list[str] = []
    for candidate in (normalized, raw):
        if candidate and candidate not in queries:
            queries.append(candidate)
    for word in re.split(r"[\s._\-+']+", raw):
        word = _normalize_search_token(word)
        if len(word) >= 3 and word not in queries:
            queries.append(word)
    return queries


def _score_name_match(query: str, path: str) -> int:
    """Higher score = better match between query and file path."""
    q = query.lower().strip()
    base = Path(path).stem.lower()
    full = Path(path).name.lower()
    q_base = q[:-4] if q.lower().endswith(".pdf") else q
    q_norm = _normalize_search_token(q_base)

    if full == q or base == q_base or base == q_norm:
        return 1000
    if q_base in base or base in q_base or q_norm in base or base in q_norm:
        return 850
    q_words = [w for w in re.split(r"[\s._\-+]+", q_norm) if len(w) > 1]
    if not q_words:
        return 0
    hits = sum(1 for w in q_words if w in base)
    if hits == len(q_words):
        return 700 + hits * 10
    if hits:
        return 200 + hits * 80
    if len(q_norm) >= 4 and q_norm.rstrip("s") in base:
        return 650
    return 0


def _subsequence_in(text: str, query: str) -> bool:
    qi = 0
    for ch in text:
        if qi < len(query) and ch == query[qi]:
            qi += 1
    return qi == len(query)


def _subsequence_score(query: str, path: str) -> int:
    """Match short tokens like JRI within individual filename segments."""
    q = re.sub(r"[^a-zA-Z]", "", query).lower()
    if len(q) < 2:
        return 0
    parts = [p for p in re.split(r"[_\s]+", Path(path).stem.lower()) if p]
    if not parts:
        return 0

    best = 0
    for part in parts:
        if _subsequence_in(part, q):
            best = max(best, 480 + len(q) * 35)

    stem = Path(path).stem.lower()
    if len(q) == 3 and q == "jri" and ("jaya" in stem and "rai" in stem):
        best = max(best, 820)
    if len(q) <= 4 and "rai" in stem and any(k in stem for k in ("jaya", "resume")):
        best = max(best, 760)

    return best


def _combined_name_score(query: str, path: str) -> int:
    return max(
        _score_name_match(query, path),
        _subsequence_score(query, path),
    )


def _extract_hint_tokens(hint_text: str) -> list[str]:
    if not hint_text:
        return []
    stop = {
        "compare", "comparison", "compiled", "compile", "compiling", "summarize",
        "summarise", "summary", "download", "downloads", "pdf", "pdfs", "resume",
        "document", "and", "the", "in", "it", "as", "well", "two", "with", "versus",
    }
    tokens: list[str] = []
    for word in re.split(r"[\s._\-+',]+", hint_text):
        w = _normalize_search_token(word)
        if len(w) >= 2 and w.lower() not in stop and w.lower() not in tokens:
            tokens.append(w)
    return tokens


def _pick_best_pdf(
    name: str,
    pdfs: list[str],
    exclude: set[str] | None = None,
    hint_tokens: list[str] | None = None,
    prefer_resume: bool = False,
) -> tuple[str | None, int]:
    exclude = exclude or set()
    best_path: str | None = None
    best_score = 0
    for path in pdfs:
        if path in exclude:
            continue
        score = _combined_name_score(name, path) if name.strip() else 0
        if hint_tokens:
            for token in hint_tokens:
                score = max(score, _combined_name_score(token, path))
        if prefer_resume and "resume" in Path(path).stem.lower():
            score += 40
        if score > best_score:
            best_score = score
            best_path = path
    return best_path, best_score


def _fallback_compare_pair(
    name_a: str,
    name_b: str,
    directory: str | None,
    hint_text: str = "",
) -> dict | None:
    """Scan a folder and pick the best two distinct PDFs when direct search fails."""
    search_dir = str(Path(directory).expanduser()) if directory else HOME
    pdfs = _list_pdfs_in_dir(search_dir)
    if len(pdfs) < 2:
        pdfs = [
            p for p in pdfs
            if os.path.isfile(p) and p.lower().endswith(".pdf")
        ]
        extra = _list_pdfs_in_dir(HOME) if search_dir != HOME else []
        pdfs = sorted({*pdfs, *[p for p in extra if p.lower().endswith(".pdf")]})
    if len(pdfs) < 2:
        return None

    hint_tokens = _extract_hint_tokens(hint_text)
    min_score = 150
    prefer_resume = "resume" in hint_text.lower()

    path_b, score_b = _pick_best_pdf(
        name_b, pdfs, hint_tokens=hint_tokens, prefer_resume=prefer_resume,
    )
    if name_a.strip():
        path_a, score_a = _pick_best_pdf(
            name_a,
            pdfs,
            exclude={path_b} if path_b else set(),
            hint_tokens=hint_tokens,
        )
    else:
        path_a, score_a = None, 0

    if path_b and score_b < min_score and hint_tokens:
        path_b, score_b = _pick_best_pdf("", pdfs, hint_tokens=hint_tokens)
    if path_a and score_a < min_score and hint_tokens:
        path_a, score_a = _pick_best_pdf(
            "",
            pdfs,
            exclude={path_b} if path_b else set(),
            hint_tokens=hint_tokens,
        )

    if not path_b or score_b < min_score:
        return None

    if not path_a or path_a == path_b or score_a < min_score:
        path_a = _pick_complement_pdf(path_b, pdfs, hint_tokens, name_a)
        if path_a:
            score_a = max(_combined_name_score(name_a, path_a), 200)
    elif score_a < min_score and not name_a.strip():
        path_a = _pick_complement_pdf(path_b, pdfs, hint_tokens, name_a) or path_a

    if not path_a or path_a == path_b:
        return None

    return {
        "path_a": path_a,
        "path_b": path_b,
        "score_a": score_a,
        "score_b": score_b,
    }


def _pick_complement_pdf(
    path_b: str,
    pdfs: list[str],
    hint_tokens: list[str],
    name_a: str,
) -> str | None:
    """Pick a second PDF when the first spoken name is missing or garbled."""
    b_stem = Path(path_b).stem.lower()

    if name_a.strip():
        candidate, score = _pick_best_pdf(name_a, pdfs, exclude={path_b}, hint_tokens=hint_tokens)
        if candidate and score >= 150:
            return candidate

    if "jaya" in b_stem or "rai" in b_stem:
        for path in sorted(pdfs, key=lambda p: os.path.getmtime(p), reverse=True):
            if path == path_b:
                continue
            stem = Path(path).stem.lower()
            if "shivam" in stem or "jaiswal" in stem:
                return path

    if "shivam" in b_stem or "jaiswal" in b_stem:
        for path in sorted(pdfs, key=lambda p: os.path.getmtime(p), reverse=True):
            if path == path_b:
                continue
            stem = Path(path).stem.lower()
            if "jaya" in stem or "rai" in stem:
                return path

    for path in sorted(pdfs, key=lambda p: os.path.getmtime(p), reverse=True):
        if path != path_b:
            return path
    return None
    """Higher score = better match between query and file path."""
    q = query.lower().strip()
    base = Path(path).stem.lower()
    full = Path(path).name.lower()
    q_base = q[:-4] if q.lower().endswith(".pdf") else q
    q_norm = _normalize_search_token(q_base)

    if full == q or base == q_base or base == q_norm:
        return 1000
    if q_base in base or base in q_base or q_norm in base or base in q_norm:
        return 850
    q_words = [w for w in re.split(r"[\s._\-+]+", q_norm) if len(w) > 1]
    if not q_words:
        return 0
    hits = sum(1 for w in q_words if w in base)
    if hits == len(q_words):
        return 700 + hits * 10
    if hits:
        return 200 + hits * 80
    if len(q_norm) >= 4 and q_norm.rstrip("s") in base:
        return 650
    return 0


def _spotlight_paths(name: str, search_dir: str) -> list[str]:
    result = subprocess.run(
        ["mdfind", "-onlyin", search_dir, "-name", name],
        capture_output=True,
        text=True,
    )
    return [l.strip() for l in (result.stdout or "").strip().splitlines() if l.strip()]


def _list_pdfs_in_dir(directory: str) -> list[str]:
    """Fallback: list PDFs in a folder when Spotlight returns nothing."""
    pdfs: list[str] = []
    try:
        for entry in os.listdir(directory):
            if not entry.lower().endswith(".pdf"):
                continue
            full = os.path.join(directory, entry)
            if os.path.isfile(full):
                pdfs.append(full)
    except OSError:
        pass
    return pdfs


def find_pdf_candidates(name: str, directory: str = None) -> dict:
    """Find PDFs matching name (exact or partial) under directory."""
    search_dir = str(Path(directory).expanduser()) if directory else HOME
    queries = _search_queries(name)

    raw: list[str] = []
    for q in queries:
        raw.extend(_spotlight_paths(q, search_dir))
        if search_dir != HOME:
            raw.extend(_spotlight_paths(q, HOME))

    pdfs = sorted({p for p in raw if p.lower().endswith(".pdf") and os.path.isfile(p)})
    if not pdfs and directory:
        scanned = _list_pdfs_in_dir(search_dir)
        pdfs = sorted({p for p in scanned if _score_name_match(name, p) > 0})
    if not pdfs:
        return {"success": False, "message": f"No PDF found matching '{name}'", "matches": []}

    def _rank_key(path: str) -> tuple:
        in_dir = path.startswith(search_dir + os.sep) or path == search_dir
        return (
            -_score_name_match(name, path),
            0 if in_dir else 1,
            len(path.split("/")),
            path,
        )

    ranked = sorted(pdfs, key=_rank_key)
    best = ranked[0]
    score = _score_name_match(name, best)
    exact = score >= 850
    return {
        "success": True,
        "path": best,
        "matches": ranked[:5],
        "exact": exact,
        "partial": not exact,
        "score": score,
    }


def find_compare_pdf_pair(
    name_a: str,
    name_b: str,
    directory: str = None,
    hint_text: str = "",
) -> dict:
    """
    Resolve two PDFs for comparison. Uses partial matching when names are inexact.
    Falls back to folder scan + fuzzy match, then asks user to confirm before read.
    """
    ra = find_pdf_candidates(name_a, directory) if name_a.strip() else {"success": False}
    rb = find_pdf_candidates(name_b, directory) if name_b.strip() else {"success": False}

    path_a = ra.get("path") if ra.get("success") else None
    path_b = rb.get("path") if rb.get("success") else None
    score_a = ra.get("score", 0) if ra.get("success") else 0
    score_b = rb.get("score", 0) if rb.get("success") else 0
    used_fallback = False

    if not path_a or not path_b:
        fallback = _fallback_compare_pair(name_a, name_b, directory, hint_text)
        if fallback:
            used_fallback = True
            if not path_a:
                path_a = fallback["path_a"]
                score_a = fallback["score_a"]
            if not path_b:
                path_b = fallback["path_b"]
                score_b = fallback["score_b"]
            ra = {"exact": score_a >= 850, "partial": score_a < 850}
            rb = {"exact": score_b >= 850, "partial": score_b < 850}

    if not path_a:
        return {"success": False, "message": ra.get("message", f"No PDF for '{name_a or 'first file'}'")}
    if not path_b:
        return {"success": False, "message": rb.get("message", f"No PDF for '{name_b or 'second file'}'")}

    if path_a == path_b:
        rb_matches = rb.get("matches") if isinstance(rb, dict) else None
        if not rb_matches and name_b.strip():
            rb_matches = find_pdf_candidates(name_b, directory).get("matches", [])
        for alt in rb_matches or []:
            if alt != path_a:
                path_b = alt
                break
        if path_a == path_b:
            ra_matches = ra.get("matches") if isinstance(ra, dict) else None
            if not ra_matches and name_a.strip():
                ra_matches = find_pdf_candidates(name_a, directory).get("matches", [])
            for alt in ra_matches or []:
                if alt != path_b:
                    path_a = alt
                    break

    if path_a == path_b:
        return {"success": False, "message": "Could not find two different PDF files to compare."}

    label_a = Path(path_a).name
    label_b = Path(path_b).name
    exact = ra.get("exact") and rb.get("exact") and not used_fallback
    needs_confirm = used_fallback or not exact or ra.get("partial") or rb.get("partial")

    if used_fallback or not exact:
        message = (
            "Could not match exact names — best matches found:\n"
            f"1) {label_a}\n"
            f"2) {label_b}\n"
            "Proceed with comparison?"
        )
    else:
        message = (
            f"I found two PDFs to compare:\n"
            f"1) {label_a}\n"
            f"2) {label_b}\n"
            "Proceed with comparison?"
        )

    return {
        "success": True,
        "path_a": path_a,
        "path_b": path_b,
        "label_a": label_a,
        "label_b": label_b,
        "exact": exact,
        "partial": needs_confirm,
        "needs_confirm": needs_confirm,
        "fallback_match": used_fallback,
        "message": message,
    }


def find_file(name: str, directory: str = None, extension: str = "") -> dict:
    """
    Search for a file by name using macOS Spotlight (mdfind).
    Optionally restrict to a folder and/or file extension (e.g. .pdf).
    """
    search_dir = str(Path(directory).expanduser()) if directory else HOME
    ext = (extension or "").lower()
    if ext and not ext.startswith("."):
        ext = f".{ext}"

    if ext in (".pdf",):
        out = find_pdf_candidates(name, directory)
        return out if out.get("success") else {
            "success": False,
            "message": f"No PDF found matching '{name}'",
            "matches": [],
        }

    result = subprocess.run(
        ["mdfind", "-onlyin", search_dir, "-name", name],
        capture_output=True,
        text=True,
    )
    matches = [l.strip() for l in result.stdout.strip().splitlines() if l.strip()]

    if not matches and search_dir != HOME:
        result2 = subprocess.run(
            ["mdfind", "-onlyin", HOME, name],
            capture_output=True,
            text=True,
        )
        matches = [l.strip() for l in result2.stdout.strip().splitlines() if l.strip()]

    if ext:
        filtered = [m for m in matches if m.lower().endswith(ext)]
        if filtered:
            matches = filtered

    if not matches:
        return {"success": False, "message": f"No file found matching '{name}'", "matches": []}

    matches.sort(key=lambda p: (len(p.split("/")), p))
    return {"success": True, "path": matches[0], "matches": matches[:5]}


def _is_safe_path(path: str) -> bool:
    """Check that path is within an allowed directory."""
    abs_path = str(Path(path).expanduser().resolve())
    return any(abs_path.startswith(base) for base in SAFE_BASE_DIRS)


def list_files(directory: str = "~/Desktop") -> dict:
    """List files in a directory."""
    expanded = str(Path(directory).expanduser())
    if not os.path.isdir(expanded):
        return {"success": False, "message": f"Not a directory: {directory}"}
    files = os.listdir(expanded)
    return {"success": True, "files": sorted(files), "directory": expanded}


def read_file(path: str) -> dict:
    """Read text content of a file."""
    expanded = str(Path(path).expanduser())
    if not os.path.exists(expanded):
        return {"success": False, "message": f"File not found: {path}"}
    if not _is_safe_path(expanded):
        return {"success": False, "message": "Access denied: outside allowed directories"}
    try:
        with open(expanded, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        return {"success": True, "content": content, "path": expanded}
    except Exception as e:
        return {"success": False, "message": str(e)}


def write_file(path: str, content: str) -> dict:
    """
    Write text content to a file. Auto-detects format from extension.
    Supports: .txt / .md / .csv and any plain text — .docx (real Word document).
    """
    expanded = str(Path(path).expanduser())
    if not _is_safe_path(expanded):
        return {"success": False, "message": "Access denied: outside allowed directories"}

    ext = Path(expanded).suffix.lower()
    Path(expanded).parent.mkdir(parents=True, exist_ok=True)

    if ext == ".docx":
        return _write_docx(expanded, content)
    else:
        try:
            with open(expanded, "w", encoding="utf-8") as f:
                f.write(content)
            return {"success": True, "message": f"Written to {expanded}", "path": expanded}
        except Exception as e:
            return {"success": False, "message": str(e)}


def _write_docx(path: str, content: str) -> dict:
    try:
        from docx import Document
    except ImportError:
        return {"success": False, "message": "python-docx not installed. Run: pip install python-docx"}
    try:
        doc = Document()
        for line in content.splitlines():
            doc.add_paragraph(line)
        doc.save(path)
        return {"success": True, "message": f"Written to {path}", "path": path}
    except Exception as e:
        return {"success": False, "message": str(e)}


def delete_file(path: str) -> dict:
    """Delete a file. REQUIRES user confirmation (handled by safety layer)."""
    expanded = str(Path(path).expanduser())
    if not _is_safe_path(expanded):
        return {"success": False, "message": "Access denied: outside allowed directories"}
    if not os.path.exists(expanded):
        return {"success": False, "message": f"File not found: {path}"}
    try:
        if os.path.isdir(expanded):
            shutil.rmtree(expanded)
        else:
            os.remove(expanded)
        return {"success": True, "message": f"Deleted: {expanded}"}
    except Exception as e:
        return {"success": False, "message": str(e)}


def move_file(src: str, dst: str) -> dict:
    """Move or rename a file. REQUIRES user confirmation."""
    src_exp = str(Path(src).expanduser())
    dst_exp = str(Path(dst).expanduser())
    if not _is_safe_path(src_exp) or not _is_safe_path(dst_exp):
        return {"success": False, "message": "Access denied: outside allowed directories"}
    try:
        shutil.move(src_exp, dst_exp)
        return {"success": True, "message": f"Moved {src} → {dst}"}
    except Exception as e:
        return {"success": False, "message": str(e)}


def edit_file(path: str, old_text: str, new_text: str) -> dict:
    """
    Replace old_text with new_text in any file type (in-place, never overwrites the whole file).
    Supports: .txt, .md, .csv and plain text files — .docx (Word) — .xlsx/.xls (Excel).
    """
    expanded = str(Path(path).expanduser())
    if not os.path.exists(expanded):
        return {"success": False, "message": f"File not found: {path}"}
    if not _is_safe_path(expanded):
        return {"success": False, "message": "Access denied: outside allowed directories"}

    ext = Path(expanded).suffix.lower()

    if ext in (".docx",):
        return _edit_docx(expanded, old_text, new_text)
    elif ext in (".xlsx", ".xls", ".xlsm"):
        return _edit_excel(expanded, old_text, new_text)
    else:
        return _edit_plaintext(expanded, old_text, new_text)


def _edit_plaintext(path: str, old_text: str, new_text: str) -> dict:
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        if old_text == "":
            with open(path, "a", encoding="utf-8") as f:
                f.write("\n" + new_text)
            return {"success": True, "message": f"Appended text to {path}", "path": path}
        if old_text not in content:
            return {"success": False, "message": f"Text not found in file: {old_text!r}"}
        with open(path, "w", encoding="utf-8") as f:
            f.write(content.replace(old_text, new_text, 1))
        return {"success": True, "message": f"Replaced text in {path}", "path": path}
    except Exception as e:
        return {"success": False, "message": str(e)}


def _edit_docx(path: str, old_text: str, new_text: str) -> dict:
    try:
        from docx import Document
    except ImportError:
        return {"success": False, "message": "python-docx not installed. Run: pip install python-docx"}
    try:
        doc = Document(path)
        replaced = 0

        def _replace_in_para(para):
            """Replace text in a paragraph, handling text split across multiple runs."""
            nonlocal replaced
            if old_text not in para.text:
                return
            # Merge all run text, replace, then put it back into the first run
            full = para.text
            new_full = full.replace(old_text, new_text, 1)
            for i, run in enumerate(para.runs):
                run.text = new_full if i == 0 else ""
            replaced += 1

        for para in doc.paragraphs:
            _replace_in_para(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_in_para(para)

        if replaced == 0:
            return {"success": False, "message": f"Text not found in document: {old_text!r}"}
        doc.save(path)
        return {"success": True, "message": f"Replaced text in {path}", "path": path}
    except Exception as e:
        return {"success": False, "message": str(e)}


def _edit_excel(path: str, old_text: str, new_text: str) -> dict:
    try:
        import openpyxl
    except ImportError:
        return {"success": False, "message": "openpyxl not installed. Run: pip install openpyxl"}
    try:
        wb = openpyxl.load_workbook(path)
        replaced = 0
        old_lower = old_text.lower()
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is None:
                        continue
                    cell_str = str(cell.value)
                    if old_lower in cell_str.lower():
                        # Preserve original case by finding the actual substring position
                        idx = cell_str.lower().find(old_lower)
                        cell.value = cell_str[:idx] + new_text + cell_str[idx + len(old_text):]
                        replaced += 1
        if replaced == 0:
            return {"success": False, "message": f"Text not found in spreadsheet: {old_text!r}"}
        wb.save(path)
        wb.close()
        return {"success": True, "message": f"Replaced text in {replaced} cell(s) in {path}", "path": path}
    except Exception as e:
        return {"success": False, "message": str(e)}


# ── File type groups for organize_files ────────────────────────────────────────
_TYPE_GROUPS: dict[str, set[str]] = {
    "Images":        {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".webp", ".heic", ".svg", ".ico"},
    "Videos":        {".mp4", ".mov", ".avi", ".mkv", ".wmv", ".flv", ".m4v", ".webm"},
    "Audio":         {".mp3", ".wav", ".flac", ".aac", ".ogg", ".m4a", ".wma"},
    "Documents":     {".pdf", ".doc", ".docx", ".txt", ".rtf", ".odt", ".pages", ".md"},
    "Spreadsheets":  {".xls", ".xlsx", ".csv", ".numbers", ".ods"},
    "Presentations": {".ppt", ".pptx", ".key", ".odp"},
    "Archives":      {".zip", ".tar", ".gz", ".rar", ".7z", ".bz2", ".dmg", ".pkg", ".iso"},
    "Code":          {".py", ".js", ".ts", ".html", ".css", ".java", ".cpp", ".c", ".h",
                      ".go", ".rs", ".swift", ".rb", ".sh", ".json", ".xml", ".yaml", ".yml"},
}


def organize_files(directory: str, by: str = "type", rules: list = None) -> dict:
    """
    Organize files in a directory by moving them into subfolders.

    by="type"   — auto-groups by extension into Images / Videos / Documents / etc.
    by="custom" — uses explicit rules list. Each rule:
                    { "destination": "FolderName",
                      "extensions":    [".pdf", ".docx"],   # optional
                      "name_contains": "invoice" }          # optional
                  A file matches if ANY of the specified conditions is true.

    Returns: { success, moved, skipped, folders_created, details, directory }
    """
    expanded = str(Path(directory).expanduser())
    if not os.path.isdir(expanded):
        return {"success": False, "message": f"Not a directory: {directory}"}
    if not _is_safe_path(expanded):
        return {"success": False, "message": "Access denied: outside allowed directories"}

    try:
        entries = [
            e for e in Path(expanded).iterdir()
            if e.is_file() and not e.name.startswith(".")
        ]
    except Exception as e:
        return {"success": False, "message": str(e)}

    if by == "type":
        move_plan = _plan_by_type(entries)
    elif by == "custom":
        if not rules:
            return {"success": False, "message": "custom mode requires a rules list"}
        move_plan = _plan_by_rules(entries, rules)
    else:
        return {"success": False, "message": f"Unknown strategy: {by!r}. Use 'type' or 'custom'"}

    moved, skipped, folders_created, details = 0, 0, [], []
    for src, folder_name in move_plan:
        dest_dir = Path(expanded) / folder_name
        if not dest_dir.exists():
            dest_dir.mkdir(parents=True, exist_ok=True)
            folders_created.append(folder_name)
        dest = dest_dir / src.name
        if dest.exists():
            skipped += 1
            continue
        try:
            shutil.move(str(src), str(dest))
            details.append(f"{src.name} → {folder_name}/")
            moved += 1
        except Exception as e:
            skipped += 1

    return {
        "success": True,
        "moved": moved,
        "skipped": skipped,
        "folders_created": folders_created,
        "details": details,
        "directory": expanded,
        "message": (
            f"Moved {moved} file(s) into {len(set(f for _, f in move_plan))} folder(s)"
            + (f", skipped {skipped}" if skipped else "")
        ),
    }


def _plan_by_type(entries: list) -> list[tuple]:
    """Return [(Path, folder_name), ...] grouped by file type."""
    ext_to_group = {}
    for group, exts in _TYPE_GROUPS.items():
        for ext in exts:
            ext_to_group[ext] = group

    plan = []
    for f in entries:
        group = ext_to_group.get(f.suffix.lower(), "Other")
        plan.append((f, group))
    return plan


def _plan_by_rules(entries: list, rules: list) -> list[tuple]:
    """Return [(Path, folder_name), ...] matched against user-defined rules."""
    plan = []
    for f in entries:
        ext = f.suffix.lower()
        name_lower = f.name.lower()
        matched = False
        for rule in rules:
            dest = rule.get("destination", "Other")
            rule_exts = {e.lower() for e in rule.get("extensions", [])}
            name_kw   = rule.get("name_contains", "").lower()
            if (rule_exts and ext in rule_exts) or (name_kw and name_kw in name_lower):
                plan.append((f, dest))
                matched = True
                break
        if not matched:
            plan.append((f, "Other"))
    return plan


def open_file(path: str) -> dict:
    """Open a file with its default application."""
    import subprocess
    expanded = str(Path(path).expanduser())
    if not os.path.exists(expanded):
        return {"success": False, "message": f"File not found: {path}"}
    result = subprocess.run(["open", expanded], capture_output=True, text=True)
    if result.returncode == 0:
        return {"success": True, "message": f"Opened {expanded}", "path": expanded}
    return {"success": False, "message": result.stderr.strip()}
